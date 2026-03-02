#!/usr/bin/env python3
from docx import Document

doc = Document('79.BCA-Ricky Ardi Suwanto.docx')

print('='*100)
print('ANALYZING: 79.BCA-Ricky Ardi Suwanto.docx')
print('='*100)

print('\n[PARAGRAPHS]')
for i, para in enumerate(doc.paragraphs[:25]):
    if para.text.strip():
        print(f'{i}: {para.text}')

print('\n[TABLES]')
for t_idx, table in enumerate(doc.tables):
    print(f'\nTable {t_idx}: {len(table.rows)} rows x {len(table.columns)} columns')
    
    if table.rows:
        print('HEADER ROW:')
        header = [c.text.strip() for c in table.rows[0].cells]
        for col_idx, col_name in enumerate(header):
            print(f'  [{col_idx}]: {col_name}')
        
        print('\nDATA ROWS:')
        for row_idx in range(1, min(3, len(table.rows))):
            row = table.rows[row_idx]
            print(f'\n  Row {row_idx}:')
            for col_idx, cell in enumerate(row.cells):
                cell_val = cell.text.strip()
                if cell_val:
                    header_name = header[col_idx] if col_idx < len(header) else 'UNKNOWN'
                    print(f'    {header_name}: {cell_val}')
